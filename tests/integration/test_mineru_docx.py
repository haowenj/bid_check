from __future__ import annotations

import zipfile
from pathlib import Path

import pytest
from docx import Document

from bid_check.mineru_client import MinerUClient
from generate_test_docx import generate_fixture


def test_generate_fixture_contains_all_structural_markers(tmp_path: Path):
    output_path = generate_fixture(tmp_path / "features.docx")
    document = Document(output_path)
    paragraphs = document.paragraphs
    paragraph_text = [paragraph.text for paragraph in paragraphs]

    assert output_path.is_file()
    assert "第三章 技术要求" in paragraph_text
    assert "3.1 总体要求" in paragraph_text
    assert "3.1.2 人员要求" in paragraph_text
    assert "项目经理应具有三年以上同类项目经验。" in paragraph_text
    assert "3.1.3 其他人员要求" in paragraph_text
    assert "第四章 商务要求" in paragraph_text
    assert "商务响应应完整。" in paragraph_text
    assert any(paragraph.style.name == "List Bullet" for paragraph in paragraphs)
    assert any(paragraph.style.name == "List Number" for paragraph in paragraphs)
    assert sum(not paragraph.text.strip() for paragraph in paragraphs) >= 2
    assert len(document.tables) == 1
    assert len(document.tables[0].rows) >= 3
    assert len(document.inline_shapes) >= 1

    with zipfile.ZipFile(output_path) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
    assert "<m:oMathPara" in xml
    assert "<m:oMath" in xml


@pytest.mark.integration
def test_live_mineru_returns_docx_office_artifacts(tmp_path: Path):
    fixture = Path(__file__).parents[1] / "fixtures/generated/mineru_docx_features.docx"
    artifacts = MinerUClient("http://127.0.0.1:7100", 1800).parse_docx(
        fixture, tmp_path / "outputs"
    )

    json_files = [path for path in artifacts.extracted_files if path.suffix == ".json"]
    json_names = {path.name for path in json_files}
    assert (artifacts.raw_dir / "response.zip").is_file()
    assert any(name.endswith("_middle.json") for name in json_names)
    assert any(name.endswith("_content_list.json") for name in json_names)
    assert any(name.endswith("_content_list_v2.json") for name in json_names)
    for path in json_files:
        assert path.read_text(encoding="utf-8")
    assert any(len(path.parts) >= 2 and path.parts[-2] == "images" for path in artifacts.extracted_files)
