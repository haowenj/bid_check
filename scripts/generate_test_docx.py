"""Generate a deterministic DOCX covering the first-stage normalization cases."""

from __future__ import annotations

import argparse
from io import BytesIO
from pathlib import Path
from typing import Sequence

from PIL import Image, ImageDraw
from docx import Document
from docx.oxml import OxmlElement


def _append_formula(document: Document) -> None:
    math_para = OxmlElement("m:oMathPara")
    math = OxmlElement("m:oMath")
    math_run = OxmlElement("m:r")
    math_text = OxmlElement("m:t")
    math_text.text = "E = mc²"
    math_run.append(math_text)
    math.append(math_run)
    math_para.append(math)

    body = document._body._element
    body.insert(len(body) - 1, math_para)


def _image_bytes() -> BytesIO:
    image = Image.new("RGB", (320, 120), "#e8f1ff")
    draw = ImageDraw.Draw(image)
    draw.rectangle((4, 4, 316, 116), outline="#174a7e", width=3)
    draw.text((24, 48), "Bid Check Fixture", fill="#174a7e")
    result = BytesIO()
    image.save(result, format="PNG")
    result.seek(0)
    return result


def generate_fixture(output_path: Path) -> Path:
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()

    document.add_heading("第三章 技术要求", level=1)
    document.add_heading("3.1 总体要求", level=2)
    document.add_heading("3.1.2 人员要求", level=3)
    document.add_paragraph("项目经理应具有三年以上同类项目经验。")
    document.add_paragraph("项目团队应满足招标文件规定的岗位要求。")
    document.add_heading("3.1.3 其他人员要求", level=3)
    document.add_paragraph("技术负责人应具有相关专业高级职称。")
    document.add_heading("第四章 商务要求", level=1)
    document.add_paragraph("商务响应应完整。")

    document.add_paragraph("第一项要求", style="List Bullet")
    document.add_paragraph("第二项要求", style="List Bullet")
    document.add_paragraph("第一步报价", style="List Number")
    document.add_paragraph("第二步承诺", style="List Number")

    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "角色"
    table.rows[0].cells[1].text = "要求"
    for role, requirement in (
        ("项目经理", "三年以上经验"),
        ("技术负责人", "高级职称"),
    ):
        cells = table.add_row().cells
        cells[0].text = role
        cells[1].text = requirement

    document.add_paragraph("人员配置表")
    document.add_picture(_image_bytes(), width=None)
    document.add_paragraph("图 1 系统架构示意图")
    document.add_paragraph("公式示例：")
    _append_formula(document)

    document.add_paragraph("")
    document.add_paragraph("   ")
    document.save(output_path)
    return output_path


def main(argv: Sequence[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args(argv)
    generate_fixture(args.output)
    print(args.output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
